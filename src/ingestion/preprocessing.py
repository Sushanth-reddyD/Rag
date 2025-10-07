"""Document preprocessing and parsing for various formats."""

import hashlib
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
from datetime import datetime

from pydantic import BaseModel


class DocumentParser:
    """Base parser for extracting text and structure from documents."""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document and extract text with structural information.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (text_content, structural_metadata)
        """
        raise NotImplementedError("Subclasses must implement parse()")
    
    def _compute_checksum(self, file_path: str) -> str:
        """Compute MD5 checksum of file."""
        md5_hash = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)
        return md5_hash.hexdigest()


class TextParser(DocumentParser):
    """Parser for plain text files."""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        metadata = {
            'source_type': 'txt',
            'checksum': self._compute_checksum(file_path),
            'file_size': Path(file_path).stat().st_size,
            'sections': []  # Plain text has no inherent structure
        }
        
        return text, metadata


class PDFParser(DocumentParser):
    """Parser for PDF files with structure preservation."""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF file, preserving page and section information."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")
        
        text_parts = []
        sections = []
        
        with open(file_path, 'rb') as f:
            pdf = pypdf.PdfReader(f)
            num_pages = len(pdf.pages)
            
        with open(file_path, 'rb') as f:
            pdf = pypdf.PdfReader(f)
            num_pages = len(pdf.pages)
            current_offset = 0

            for page_num, page in enumerate(pdf.pages):
                # Guard against None from extract_text()
                page_text = page.extract_text() or ""
                text_parts.append(page_text)

                page_length = len(page_text)
                sections.append({
                    'type': 'page',
                    'number': page_num + 1,
                    'start_offset': current_offset,
                    'end_offset': current_offset + page_length
                })
                current_offset += page_length

        full_text = "\n\n".join(text_parts)

        metadata = {
            'source_type': 'pdf',
            'checksum': self._compute_checksum(file_path),
            'file_size': Path(file_path).stat().st_size,
            'num_pages': num_pages,
            'sections': sections
        }
        
        return full_text, metadata


class HTMLParser(DocumentParser):
    """Parser for HTML files with semantic structure."""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse HTML file, extracting semantic structure."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required for HTML parsing. Install with: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title
        title = soup.title.string if soup.title else None
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text(separator='\n', strip=True)
        
        # Extract structure (headings, paragraphs)
        sections = []
        current_offset = 0
        
        sections = []
        search_pos = 0

        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            element_text = element.get_text(strip=True)
            if not element_text:
                continue

            start_idx = text.find(element_text, search_pos)
            if start_idx == -1:
                continue

            end_idx = start_idx + len(element_text)
            sections.append({
                'type': 'heading',
                'level': element.name,
                'text': element_text,
                'start_offset': start_idx,
                'end_offset': end_idx
            })
            search_pos = end_idx

        metadata = {
            'source_type': 'html',
            'checksum': self._compute_checksum(file_path),
            'file_size': Path(file_path).stat().st_size,
            'title': title,
            'sections': sections
        }
        
        return text, metadata


class DocxParser(DocumentParser):
    """Parser for DOCX files."""
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file, preserving paragraph structure."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        doc = docx.Document(file_path)
        
        text_parts = []
        sections = []
        current_offset = 0
        
        for para in doc.paragraphs:
            para_text = para.text
            if para_text.strip():
                text_parts.append(para_text)
                
                # Check if paragraph is a heading
                is_heading = para.style.name.startswith('Heading')
                
                sections.append({
                    'type': 'heading' if is_heading else 'paragraph',
                    'style': para.style.name,
                    'start_offset': current_offset,
                    'end_offset': current_offset + len(para_text)
                })
                
                current_offset += len(para_text) + 1  # +1 for newline
        
        full_text = "\n".join(text_parts)
        
        metadata = {
            'source_type': 'docx',
            'checksum': self._compute_checksum(file_path),
            'file_size': Path(file_path).stat().st_size,
            'num_paragraphs': len(doc.paragraphs),
            'sections': sections
        }
        
        return full_text, metadata


class DocumentPreprocessor:
    """Main preprocessor that handles multiple document formats."""
    
    PARSERS = {
        '.txt': TextParser,
        '.pdf': PDFParser,
        '.html': HTMLParser,
        '.htm': HTMLParser,
        '.docx': DocxParser,
    }
    
    def __init__(self):
        """Initialize preprocessor with available parsers."""
        self.parsers = {
            ext: parser_class() for ext, parser_class in self.PARSERS.items()
        }
    
    def process_document(
        self,
        file_path: str,
        document_id: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Process a document file and extract text with metadata.
        
        Args:
            file_path: Path to document file
            document_id: Optional document ID (generated if not provided)
            
        Returns:
            Tuple of (text_content, metadata_dict)
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Get appropriate parser
        extension = path.suffix.lower()
        parser = self.parsers.get(extension)
        
        if parser is None:
            # Fallback to text parser for unknown types
            parser = TextParser()
        
        # Parse document
        text, structural_metadata = parser.parse(str(path))
        
        # Generate document ID if not provided
        if document_id is None:
            document_id = hashlib.md5(
                f"{path.name}{structural_metadata.get('checksum', '')}".encode()
            ).hexdigest()
        
        # Build complete metadata
        metadata = {
            'document_id': document_id,
            'source': str(path.absolute()),
            'title': path.stem,
            'ingestion_timestamp': datetime.now(),
            **structural_metadata
        }
        
        return text, metadata
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        cleaned = '\n'.join(line for line in lines if line)
        
        # Normalize unicode
        cleaned = cleaned.encode('utf-8', errors='ignore').decode('utf-8')
        
        return cleaned
